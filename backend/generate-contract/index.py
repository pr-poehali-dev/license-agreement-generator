'''
Business: Generate license agreement documents package from template
Args: event with httpMethod, body containing form data
Returns: ZIP archive with 5 DOCX files or error message
'''

import json
import io
import zipfile
import urllib.request
import os
import base64
from typing import Dict, Any
from datetime import datetime

def replace_text_in_paragraph(paragraph, replacements):
    """Replace all placeholders in paragraph while preserving formatting"""
    full_text = paragraph.text
    
    # Check if any replacement is needed
    needs_replacement = False
    for key in replacements.keys():
        if key in full_text:
            needs_replacement = True
            break
    
    if not needs_replacement:
        return
    
    # Perform all replacements
    new_text = full_text
    for key, value in replacements.items():
        new_text = new_text.replace(key, value)
    
    # If text hasn't changed, skip
    if new_text == full_text:
        return
    
    # Clear existing runs and add new text with first run's formatting
    if paragraph.runs:
        first_run = paragraph.runs[0]
        # Save formatting properties
        font_name = first_run.font.name
        font_size = first_run.font.size
        bold = first_run.font.bold
        italic = first_run.font.italic
        
        # Clear all runs
        for run in paragraph.runs:
            run.text = ''
        
        # Set new text to first run
        paragraph.runs[0].text = new_text
        paragraph.runs[0].font.name = font_name
        paragraph.runs[0].font.size = font_size
        paragraph.runs[0].font.bold = bold
        paragraph.runs[0].font.italic = italic
    else:
        paragraph.text = new_text

def get_next_contract_number() -> str:
    """Get next contract number from database and increment counter"""
    import psycopg2
    
    database_url = os.environ.get('DATABASE_URL')
    if not database_url:
        raise ValueError('DATABASE_URL not found in environment')
    
    conn = psycopg2.connect(database_url)
    cur = conn.cursor()
    
    # Get and increment counter atomically
    cur.execute(
        "UPDATE contract_counter SET current_number = current_number + 1, "
        "updated_at = CURRENT_TIMESTAMP WHERE id = 1 RETURNING current_number"
    )
    new_number = cur.fetchone()[0]
    
    conn.commit()
    cur.close()
    conn.close()
    
    # Format: ЛД-ГП-001/2024
    current_year = datetime.now().year
    return f'ЛД-ГП-{new_number:03d}/{current_year}'

def send_document_to_telegram(file_bytes: bytes, filename: str, bot_token: str, chat_id: str):
    """Send single document to Telegram"""
    url = f'https://api.telegram.org/bot{bot_token}/sendDocument'
    
    boundary = '----WebKitFormBoundary7MA4YWxkTrZu0gW'
    
    body_parts = []
    body_parts.append(f'--{boundary}'.encode())
    body_parts.append(f'Content-Disposition: form-data; name="chat_id"\r\n\r\n{chat_id}'.encode())
    body_parts.append(f'--{boundary}'.encode())
    body_parts.append(f'Content-Disposition: form-data; name="document"; filename="{filename}"\r\nContent-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n'.encode())
    body_parts.append(file_bytes)
    body_parts.append(f'\r\n--{boundary}--\r\n'.encode())
    
    body = b'\r\n'.join(body_parts)
    
    headers = {
        'Content-Type': f'multipart/form-data; boundary={boundary}'
    }
    
    req = urllib.request.Request(url, data=body, headers=headers)
    urllib.request.urlopen(req)

def send_documents_to_telegram(documents: list, contract_number: str):
    """Send all documents to Telegram"""
    bot_token = os.environ.get('TELEGRAM_BOT_TOKEN')
    chat_id = os.environ.get('TELEGRAM_CHAT_ID')
    
    if not bot_token or not chat_id:
        return
    
    # Send intro message
    url = f'https://api.telegram.org/bot{bot_token}/sendMessage'
    message_data = {
        'chat_id': chat_id,
        'text': f'📄 Договор {contract_number}\n\nПакет документов:'
    }
    import json
    req = urllib.request.Request(
        url,
        data=json.dumps(message_data).encode('utf-8'),
        headers={'Content-Type': 'application/json'}
    )
    urllib.request.urlopen(req)
    
    # Send each document
    for doc_bytes, doc_name in documents:
        send_document_to_telegram(doc_bytes, doc_name, bot_token, chat_id)

def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    method: str = event.get('httpMethod', 'GET')
    
    if method == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Max-Age': '86400'
            },
            'body': '',
            'isBase64Encoded': False
        }
    
    if method != 'POST':
        return {
            'statusCode': 405,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps({'error': 'Method not allowed'}),
            'isBase64Encoded': False
        }
    
    body_data = json.loads(event.get('body', '{}'))
    
    # Generate next contract number automatically
    contract_number = get_next_contract_number()
    contract_date = body_data.get('contractDate', '')
    citizenship = body_data.get('citizenship', '')
    full_name = body_data.get('fullName', '')
    short_name = body_data.get('shortName', '')
    nickname = body_data.get('nickname', '')
    passport = body_data.get('passport', '')
    email = body_data.get('email', '')
    
    if not all([contract_date, citizenship, full_name, short_name, nickname, passport, email]):
        return {
            'statusCode': 400,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps({'error': 'All fields are required'}),
            'isBase64Encoded': False
        }
    
    template_url = 'https://disk.yandex.ru/d/iXT9pWLUhONW3A'
    
    direct_download_url = get_direct_download_link(template_url)
    
    template_response = urllib.request.urlopen(direct_download_url)
    template_bytes = template_response.read()
    
    from docx import Document
    
    replacements = {
        '{{номер_договора}}': contract_number,
        '{{дата_заключения_договора}}': contract_date,
        '{{graj}}': citizenship,
        '{{ФИО_ИП_полностью_кого}}': full_name,
        '{{ФИО_ИП_кратко}}': short_name,
        '{{NIK}}': nickname,
        '{{PAS}}': passport,
        '{{mail}}': email
    }
    
    doc_names = [
        f'Договор_{contract_number.replace("/", "-")}.docx',
        f'Приложение_1_{contract_number.replace("/", "-")}.docx',
        f'Приложение_2_{contract_number.replace("/", "-")}.docx',
        f'Приложение_3_{contract_number.replace("/", "-")}.docx',
        f'Акт_{contract_number.replace("/", "-")}.docx'
    ]
    
    documents_for_telegram = []
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for doc_name in doc_names:
            doc = Document(io.BytesIO(template_bytes))
            
            # Replace in main document paragraphs
            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
            
            # Replace in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, replacements)
            
            # Replace in headers and footers
            for section in doc.sections:
                for header in [section.header, section.footer]:
                    for paragraph in header.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    replace_text_in_paragraph(paragraph, replacements)
            
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            doc_bytes = doc_buffer.read()
            
            # Save for Telegram
            documents_for_telegram.append((doc_bytes, doc_name))
            
            # Add to ZIP for download
            zip_file.writestr(doc_name, doc_bytes)
    
    # Send to Telegram
    try:
        send_documents_to_telegram(documents_for_telegram, contract_number)
    except Exception as e:
        pass  # Continue even if Telegram fails
    
    zip_buffer.seek(0)
    zip_bytes = zip_buffer.read()
    zip_base64 = base64.b64encode(zip_bytes).decode('utf-8')
    
    return {
        'statusCode': 200,
        'headers': {
            'Content-Type': 'application/zip',
            'Content-Disposition': f'attachment; filename="Договор_пакет_{contract_number.replace("/", "-")}.zip"',
            'Access-Control-Allow-Origin': '*'
        },
        'body': zip_base64,
        'isBase64Encoded': True
    }

def get_direct_download_link(public_url: str) -> str:
    api_url = 'https://cloud-api.yandex.net/v1/disk/public/resources/download'
    params = {'public_key': public_url}
    
    import urllib.parse
    query_string = urllib.parse.urlencode(params)
    full_url = f'{api_url}?{query_string}'
    
    response = urllib.request.urlopen(full_url)
    data = json.loads(response.read().decode('utf-8'))
    
    return data['href']